import streamlit as st
import plotly.graph_objects as go
import pandas as pd
import json
import os
import io
import docx

from core.extractor import extract_text
from core.scanner import scan_document
from core.tfidf_engine import run_tfidf, get_document_topic, get_risk_score
from db.db_manager import (
    init_db, get_keywords_by_category, add_keyword, update_keyword, delete_keyword, delete_keywords_bulk,
    get_all_heading_rules, add_heading_rule, update_heading_rule, delete_heading_rule,
    get_all_category_signals, add_category_signal, update_category_signal, delete_category_signal
)
from llm.ollama_client import check_ollama_available, list_available_models, analyze_document

# Page config
st.set_page_config(
    page_title="DocScan — Secure Document Analyzer",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Completely Offline, Premium UI CSS (No Google Fonts, sleek design)
st.markdown("""
<style>
/* Stop Streamlit from dimming/blinking the screen while running */
.stApp[data-test-script-state="running"] .main { opacity: 1 !important; transition: none !important; }
div[data-testid="stAppViewBlockContainer"] { opacity: 1 !important; transition: none !important; }

/* Offline system fonts */
html, body, [class*="css"] {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
}
[data-testid="stSidebar"] {
    background: #111827; 
    border-right: 1px solid #1f2937;
}
[data-testid="stSidebar"] * {
    color: #9ca3af;
}

/* Beautiful Metric Cards */
.metric-card {
    background: #1f2937; border: 1px solid #374151;
    border-radius: 12px; padding: 1.2rem; text-align: center;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    transition: transform 0.2s;
}
.metric-card:hover { transform: translateY(-2px); }
.metric-card .value { font-size: 2.2rem; font-weight: 700; color: #60a5fa; line-height: 1.2; }
.metric-card .label { font-size: 0.8rem; color: #9ca3af; text-transform: uppercase; font-weight: 600; margin-top: 0.5rem; letter-spacing: 0.05em; }

/* Context Snippets */
.context-snippet {
    background: #1f2937; border-left: 4px solid #3b82f6;
    border-radius: 4px 8px 8px 4px; padding: 0.75rem 1rem; margin: 0.5rem 0;
    font-size: 0.85rem; font-family: 'Consolas', 'Courier New', monospace;
    color: #d1d5db; line-height: 1.6;
}
.context-snippet .hit {
    color: #fca5a5; font-weight: bold; background: rgba(248,113,113,0.1); padding: 2px 4px; border-radius: 4px;
}

/* Badges */
.badge { padding: 4px 12px; border-radius: 9999px; font-size: 0.75rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.05em; display: inline-block; }
.badge-identity    { background:#7f1d1d; color:#fca5a5; border:1px solid #991b1b; }
.badge-financial   { background:#14532d; color:#86efac; border:1px solid #166534; }
.badge-credential  { background:#4c1d95; color:#d8b4fe; border:1px solid #5b21b6; }
.badge-classification { background:#1e3a8a; color:#93c5fd; border:1px solid #1e40af; }
.badge-personal    { background:#78350f; color:#fcd34d; border:1px solid #92400e; }
.badge-health      { background:#134e4a; color:#5eead4; border:1px solid #115e59; }
.badge-technical   { background:#312e81; color:#a5b4fc; border:1px solid #3730a3; }
.badge-general     { background:#374151; color:#d1d5db; border:1px solid #4b5563; }

.risk-HIGH    { color:#ef4444 !important; font-weight:800; }
.risk-MEDIUM  { color:#f59e0b !important; font-weight:800; }
.risk-LOW     { color:#10b981 !important; font-weight:800; }
.risk-UNKNOWN { color:#9ca3af !important; font-weight:800; }

.ollama-status { padding: 8px 12px; border-radius: 8px; font-weight: 600; text-align: center; margin-bottom: 1rem; }
.ollama-online { background: rgba(16,185,129,0.1); color: #10b981; border: 1px solid #10b981; }
.ollama-offline { background: rgba(239,68,68,0.1); color: #ef4444; border: 1px solid #ef4444; }

/* Custom Streamlit adjustments */
#MainMenu {visibility:hidden;} footer {visibility:hidden;}
</style>
""", unsafe_allow_html=True)

# Initialize Database
init_db()

# State Management
if 'processed_files' not in st.session_state:
    st.session_state.processed_files = {}  # {group_name: {filename: {results, metadata, text, etc}}}
if 'staged_folders' not in st.session_state:
    st.session_state.staged_folders = {}
if 'view_mode' not in st.session_state:
    st.session_state.view_mode = 'list' # 'list', 'report', 'print'
if 'selected_file' not in st.session_state:
    st.session_state.selected_file = None
if 'selected_group' not in st.session_state:
    st.session_state.selected_group = None

if 'undo_kw' not in st.session_state: st.session_state.undo_kw = None
if 'undo_hr' not in st.session_state: st.session_state.undo_hr = None
if 'undo_cs' not in st.session_state: st.session_state.undo_cs = None

SUPPORTED_TYPES = ['.pdf', '.docx', '.xlsx', '.pptx', 'pdf', 'docx', 'xlsx', 'pptx']

# ---------------------------------------------------------------------------
# Sidebar Configuration
# ---------------------------------------------------------------------------
with st.sidebar:
    st.markdown("<h1 style='color:#f3f4f6; margin-bottom:0;'>🔍 DocScan</h1>", unsafe_allow_html=True)
    st.markdown("<p style='margin-top:0;'>Advanced Local Security Analysis</p>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # AI Backend
    ollama_ok = check_ollama_available()
    if ollama_ok:
        st.markdown("<div class='ollama-status ollama-online'>● AI Engine Online</div>", unsafe_allow_html=True)
        models = list_available_models()
        if models:
            default_idx = models.index('llama3.2') if 'llama3.2' in models else 0
            st.selectbox("Select Model", models, index=default_idx, key="ollama_model")
    else:
        st.markdown("<div class='ollama-status ollama-offline'>● AI Engine Offline</div>", unsafe_allow_html=True)
        st.caption("Install Ollama & pull 'llama3.2' for AI Insights.")

    st.markdown("---")
    
    # Keyword Management
    st.markdown("<h3 style='color:#e5e7eb;'>🔑 Keywords Management</h3>", unsafe_allow_html=True)
    
    show_kw = st.toggle("Show / Hide Keyword Settings", value=False)
    if show_kw:
        if st.session_state.undo_kw:
            deleted_words = [kw['keyword'] for kw in st.session_state.undo_kw['kws']]
            cat = st.session_state.undo_kw['category']
            st.warning(f"Deleted {len(deleted_words)} keywords from '{cat}': {', '.join(deleted_words)}")
            if st.button("↩️ Undo", key="undo_kw_main"):
                for kw in st.session_state.undo_kw['kws']:
                    add_keyword(kw['keyword'], kw['category'], kw['weight'])
                st.session_state.undo_kw = None
                st.rerun()

        kws_by_cat = get_keywords_by_category()
        
        for cat, kws in kws_by_cat.items():
            with st.expander(f"📁 {cat.capitalize()} ({len(kws)})"):

                with st.form(f"form_kw_{cat}"):
                    selected_kws = []
                    for kw in kws:
                        if st.checkbox(f"{kw['keyword']} (w: {kw['weight']})", key=f"kw_{kw['id']}"):
                            selected_kws.append(kw)
                    
                    if st.form_submit_button("🗑️ Delete Selected"):
                        if selected_kws:
                            st.session_state.undo_kw = {'category': cat, 'kws': selected_kws}
                            delete_keywords_bulk([kw['id'] for kw in selected_kws])
                            st.rerun()
                
                st.markdown("---")
                # Edit existing
                st.markdown("**Edit Keyword**")
                edit_opts = {kw['keyword']: kw for kw in kws}
                kw_to_edit = st.selectbox("Select keyword to edit", list(edit_opts.keys()), key=f"edit_sel_{cat}")
                if kw_to_edit:
                    target = edit_opts[kw_to_edit]
                    new_kw_val = st.text_input("Keyword", value=target['keyword'], key=f"ek_{target['id']}")
                    new_wt_val = st.slider("Weight", 0.5, 5.0, float(target['weight']), 0.5, key=f"ew_{target['id']}")
                    if st.button("Update Keyword", key=f"upd_{target['id']}"):
                        update_keyword(target['id'], new_kw_val, target['category'], new_wt_val)
                        st.rerun()

        with st.expander("➕ Add New Keyword"):
            existing_cats = list(get_keywords_by_category().keys())
            all_cats = sorted(list(set(existing_cats + ['general','identity','financial','credential','classification','personal','health','technical'])))
            
            with st.form("add_kw_form"):
                new_kw = st.text_input("Keyword / phrase")
                new_cat_sel = st.selectbox("Select Existing Category", ["-- Select --"] + all_cats)
                new_cat_custom = st.text_input("OR Create New Category (Overrides selection)")
                new_wt = st.slider("Risk Weight", 0.5, 5.0, 1.5, 0.5)
                if st.form_submit_button("Save Keyword"):
                    final_cat = new_cat_custom.strip().lower() if new_cat_custom.strip() else new_cat_sel
                    if new_kw.strip() and final_cat and final_cat != "-- Select --":
                        add_keyword(new_kw.strip().lower(), final_cat, new_wt)
                        st.rerun()

    st.markdown("---")

    # Advanced Rules Management
    st.markdown("<h3 style='color:#e5e7eb;'>⚙️ Advanced Rules</h3>", unsafe_allow_html=True)
    
    with st.expander("📄 Heading Rules"):
        if st.session_state.undo_hr:
            deleted_labels = [hr['label'] for hr in st.session_state.undo_hr]
            st.warning(f"Deleted {len(deleted_labels)} rules: {', '.join(deleted_labels)}")
            if st.button("↩️ Undo", key="undo_hr_btn"):
                for hr in st.session_state.undo_hr:
                    add_heading_rule(hr['pattern'], hr['label'])
                st.session_state.undo_hr = None
                st.rerun()

        with st.popover("ℹ️ Regex Help (View Rules)"):
            st.markdown("""
            **Regex Cheat Sheet:**
            - `\\b` : Word boundary (ensures it matches a whole word, e.g. `\\binvoice\\b` matches "invoice" but not "invoices")
            - `\\s+` : Matches one or more spaces (e.g. `tax\\s+invoice`)
            - `[\\s\\-]*` : Matches optional spaces or hyphens (e.g. `non[\\s\\-]*disclosure`)
            - `(?i)` : Makes the entire pattern case-insensitive
            - `|` : OR operator (e.g. `resume|cv`)
            """)
        
        rules = get_all_heading_rules()
        with st.form("form_hr"):
            selected_hrs = []
            for r in rules:
                if st.checkbox(f"**{r['label']}** (`{r['pattern']}`)", key=f"hr_{r['id']}"):
                    selected_hrs.append(r)
            
            if st.form_submit_button("🗑️ Delete Selected"):
                if selected_hrs:
                    st.session_state.undo_hr = selected_hrs
                    for r in selected_hrs:
                        delete_heading_rule(r['id'])
                    st.rerun()
        st.markdown("---")
        with st.form("add_hr_form"):
            r_label = st.text_input("Document Label (e.g., Invoice)")
            r_pattern = st.text_input("Regex Pattern (e.g., \\binvoice\\b)")
            if st.form_submit_button("Add Rule"):
                if r_label and r_pattern:
                    add_heading_rule(r_pattern, r_label)
                    st.rerun()

    with st.expander("📊 Category Signals"):
        if st.session_state.undo_cs:
            deleted_cats = [cs['category'] for cs in st.session_state.undo_cs]
            st.warning(f"Deleted {len(deleted_cats)} signals: {', '.join(deleted_cats)}")
            if st.button("↩️ Undo", key="undo_cs_btn"):
                for cs in st.session_state.undo_cs:
                    add_category_signal(cs['category'], cs['words_list'])
                st.session_state.undo_cs = None
                st.rerun()

        sigs = get_all_category_signals()
        with st.form("form_cs"):
            selected_css = []
            for s in sigs:
                if st.checkbox(f"**{s['category']}** (`{', '.join(s['words_list'][:5])}...`)", key=f"cs_{s['id']}"):
                    selected_css.append(s)
            
            if st.form_submit_button("🗑️ Delete Selected"):
                if selected_css:
                    st.session_state.undo_cs = selected_css
                    for s in selected_css:
                        delete_category_signal(s['id'])
                    st.rerun()
        st.markdown("---")
        with st.form("add_cs_form"):
            c_label = st.text_input("Category Name")
            c_words = st.text_area("Signal Words (comma separated)")
            if st.form_submit_button("Add Signals"):
                if c_label and c_words:
                    word_list = [w.strip().lower() for w in c_words.split(',')]
                    add_category_signal(c_label, word_list)
                    st.rerun()

# ---------------------------------------------------------------------------
# Core Processing Function (Runs quietly to avoid screen blinking)
# ---------------------------------------------------------------------------
def process_documents(files_dict, group_name, do_pii, do_ner, do_fuzzy, run_ai):
    """
    files_dict: {filename: file_bytes}
    """
    all_kws = []
    for cat_kws in get_keywords_by_category().values():
        all_kws.extend(cat_kws)
        
    if group_name not in st.session_state.processed_files:
        st.session_state.processed_files[group_name] = {}
        
    temp_texts = {}
    temp_results = {}
    temp_meta = {}
    
    for fname, fbytes in files_dict.items():
        extraction = extract_text(fbytes, fname)
        if extraction['error']:
            continue
            
        kw_results = scan_document(
            extraction['text'], all_kws,
            include_pii=do_pii, include_ner=do_ner, fuzzy_mode=do_fuzzy
        )
        temp_texts[fname] = extraction['text']
        temp_meta[fname] = {
            'page_count': extraction['page_count'],
            'word_count': extraction['word_count'],
            'method': extraction['method'],
            'author': extraction.get('metadata', {}).get('author', 'Unknown'),
            'title': extraction.get('metadata', {}).get('title', 'Unknown'),
        }
        temp_results[fname] = kw_results
        
    tfidf_res = run_tfidf(temp_texts)
    
    for fname in temp_results.keys():
        doc_type = get_document_topic(tfidf_res, fname, temp_texts[fname])
        risk_score, risk_lvl = get_risk_score(temp_results[fname])
        
        ai_analysis = None
        if run_ai and ollama_ok:
            ai_analysis = analyze_document(
                temp_texts[fname], temp_results[fname], tfidf_res.get(fname, []),
                model=st.session_state.get('ollama_model', 'llama3.2')
            )
            if ai_analysis:
                risk_lvl = ai_analysis.get('risk_level', risk_lvl)
                doc_type = ai_analysis.get('document_type', doc_type)
        
        st.session_state.processed_files[group_name][fname] = {
            'metadata': temp_meta[fname],
            'results': temp_results[fname],
            'tfidf': tfidf_res.get(fname, []),
            'doc_type': doc_type,
            'risk_level': risk_lvl,
            'risk_score': risk_score,
            'ai': ai_analysis,
            'text': temp_texts[fname]
        }

# ---------------------------------------------------------------------------
# Main Page Navigation & Upload
# ---------------------------------------------------------------------------
if st.session_state.view_mode == 'list':
    st.title("Secure Document Analyzer")
    st.markdown("Upload multiple files or select a local folder to scan recursively. Analysis runs 100% offline.")
    
    t_file, t_folder, t_bulk = st.tabs(["📄 Upload Files", "📥 Queue Folders", "📁 Bulk Parent Folder"])
    
    with t_file:
        col_up, col_opt = st.columns([3, 1])
        with col_up:
            uploaded_files = st.file_uploader("Drop files here", type=SUPPORTED_TYPES, accept_multiple_files=True)
        with col_opt:
            st.markdown("**Scan Options**")
            do_pii = st.checkbox("Auto PII", value=True, key="p1")
            do_ner = st.checkbox("Auto NER", value=True, key="n1")
            do_fuzzy = st.checkbox("Fuzzy Match", value=False, key="f1")
            run_ai = st.checkbox("Run AI Analysis", value=False, key="a1") if ollama_ok else False
            scan_btn = st.button("🚀 Process Files", type="primary", key="btn1")
            
        if scan_btn and uploaded_files:
            files_dict = {f.name: f.read() for f in uploaded_files}
            with st.spinner("Analyzing Documents (this may take a moment)..."):
                process_documents(files_dict, "Uploaded Files", do_pii, do_ner, do_fuzzy, run_ai)
            st.rerun()

    with t_folder:
        col_fup, col_fopt = st.columns([3, 1])
        with col_fup:
            with st.form("stage_folder_form", clear_on_submit=True):
                folder_name_input = st.text_input("Name for this folder group (e.g. Finance_Docs)")
                folder_uploaded_files = st.file_uploader("Drop folder here (or select multiple files)", type=SUPPORTED_TYPES, accept_multiple_files=True, key="folder_drop")
                add_folder_btn = st.form_submit_button("➕ Add Folder to Queue")
                
                if add_folder_btn and folder_uploaded_files:
                    folder_name = folder_name_input.strip() or f"Folder {len(st.session_state.staged_folders) + 1}"
                    files_dict = {f.name: f.read() for f in folder_uploaded_files}
                    st.session_state.staged_folders[folder_name] = files_dict
                    st.success(f"Added '{folder_name}' ({len(files_dict)} files) to queue.")
            
            if st.session_state.staged_folders:
                st.markdown("### 📥 Queued Folders")
                for fn, fdict in st.session_state.staged_folders.items():
                    st.markdown(f"- **{fn}** ({len(fdict)} files)")
                if st.button("Clear Queue"):
                    st.session_state.staged_folders = {}
                    st.rerun()

        with col_fopt:
            st.markdown("**Scan Options**")
            f_do_pii = st.checkbox("Auto PII", value=True, key="p2")
            f_do_ner = st.checkbox("Auto NER", value=True, key="n2")
            f_do_fuzzy = st.checkbox("Fuzzy Match", value=False, key="f2")
            f_run_ai = st.checkbox("Run AI Analysis", value=False, key="a2") if ollama_ok else False
            scan_folder_btn = st.button("🚀 Process All Folders", type="primary", key="btn2")
            
        if scan_folder_btn and st.session_state.staged_folders:
            with st.spinner(f"Analyzing {len(st.session_state.staged_folders)} Folders (this may take a moment)..."):
                for folder_name, files_dict in st.session_state.staged_folders.items():
                    process_documents(files_dict, folder_name, f_do_pii, f_do_ner, f_do_fuzzy, f_run_ai)
            st.session_state.staged_folders = {}
            st.rerun()

    with t_bulk:
        col_bup, col_bopt = st.columns([3, 1])
        with col_bup:
            st.markdown("If you have 10-20 folders, simply paste the path to their **Parent Folder**. The system will instantly find all sub-folders and separate them for you.")
            parent_path = st.text_input("Enter Parent Folder Path (e.g. C:/Users/Docs/All_Projects)")
        with col_bopt:
            st.markdown("**Scan Options**")
            b_do_pii = st.checkbox("Auto PII", value=True, key="p3")
            b_do_ner = st.checkbox("Auto NER", value=True, key="n3")
            b_do_fuzzy = st.checkbox("Fuzzy Match", value=False, key="f3")
            b_run_ai = st.checkbox("Run AI Analysis", value=False, key="a3") if ollama_ok else False
            scan_bulk_btn = st.button("🚀 Process Parent Folder", type="primary", key="btn3")
            
        if scan_bulk_btn and parent_path:
            if os.path.isdir(parent_path):
                found_groups = 0
                with st.spinner(f"Scanning parent folder '{parent_path}' for sub-folders..."):
                    # Get immediate subdirectories
                    for item in os.listdir(parent_path):
                        item_path = os.path.join(parent_path, item)
                        if os.path.isdir(item_path):
                            files_dict = {}
                            # Read files in this sub-folder (non-recursive to keep it strictly grouped, or recursive within it)
                            for root, dirs, files in os.walk(item_path):
                                for file in files:
                                    ext = os.path.splitext(file)[1].lower()
                                    if ext in ['.pdf', '.docx', '.xlsx', '.pptx']:
                                        fpath = os.path.join(root, file)
                                        try:
                                            with open(fpath, 'rb') as f:
                                                files_dict[file] = f.read()
                                        except Exception:
                                            pass
                            if files_dict:
                                process_documents(files_dict, item, b_do_pii, b_do_ner, b_do_fuzzy, b_run_ai)
                                found_groups += 1
                if found_groups > 0:
                    st.success(f"Processed {found_groups} sub-folders successfully!")
                    st.rerun()
                else:
                    st.warning("No supported documents found in any sub-folders.")
            else:
                st.error("Invalid parent folder path.")

    # Display List of Processed Files Grouped by Folder/Upload
    if st.session_state.processed_files:
        st.markdown("---")
        
        # Bulk Actions
        c_act1, c_act2 = st.columns([1, 4])
        with c_act1:
            if st.button("🗑️ Delete Selected"):
                for g_name in list(st.session_state.processed_files.keys()):
                    for fname in list(st.session_state.processed_files[g_name].keys()):
                        if st.session_state.get(f"sel_{g_name}_{fname}", False):
                            del st.session_state.processed_files[g_name][fname]
                    if not st.session_state.processed_files[g_name]:
                        del st.session_state.processed_files[g_name]
                st.rerun()
                
        for group_name, files in st.session_state.processed_files.items():
            if not files: continue
            
            with st.expander(f"📁 {group_name} ({len(files)} files)", expanded=True):
                for fname, data in files.items():
                    col1, col2, col3, col4, col5 = st.columns([0.5, 3, 2, 2, 2.5])
                    
                    with col1:
                        st.checkbox("Select", key=f"sel_{group_name}_{fname}", label_visibility="collapsed")
                    with col2:
                        st.markdown(f"**📄 {fname}**")
                        st.caption(f"Author: {data['metadata']['author']} | Pages: {data['metadata']['page_count']}")
                    with col3:
                        st.markdown(f"**Type:** {data['doc_type']}")
                    with col4:
                        r_col = f"risk-{data['risk_level']}"
                        st.markdown(f"<div style='margin-top:8px;'><span class='{r_col}'>{data['risk_level']} RISK</span></div>", unsafe_allow_html=True)
                    with col5:
                        c1, c2, c3 = st.columns(3)
                        if c1.button("📊 View", key=f"view_{group_name}_{fname}"):
                            st.session_state.selected_file = fname
                            st.session_state.selected_group = group_name
                            st.session_state.view_mode = 'report'
                            st.rerun()
                        if c2.button("🖨️ Print", key=f"print_{group_name}_{fname}"):
                            st.session_state.selected_file = fname
                            st.session_state.selected_group = group_name
                            st.session_state.view_mode = 'print'
                            st.rerun()
                        if c3.button("🗑️ Del", key=f"rm_{group_name}_{fname}"):
                            del st.session_state.processed_files[group_name][fname]
                            if not st.session_state.processed_files[group_name]:
                                del st.session_state.processed_files[group_name]
                            st.rerun()
                    st.markdown("<hr style='margin: 0.2rem 0; border-color:#374151;'/>", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# View Mode: Detailed Report
# ---------------------------------------------------------------------------
elif st.session_state.view_mode == 'report':
    fname = st.session_state.selected_file
    gname = st.session_state.selected_group
    data = st.session_state.processed_files[gname][fname]
    
    if st.button("← Back to Document List"):
        st.session_state.view_mode = 'list'
        st.rerun()
        
    st.title(f"📄 Analysis Report: {fname}")
    
    # Metrics
    m1, m2, m3, m4 = st.columns(4)
    m1.markdown(f'<div class="metric-card"><div class="value">{data["metadata"]["page_count"]}</div><div class="label">Pages</div></div>', unsafe_allow_html=True)
    m2.markdown(f'<div class="metric-card"><div class="value">{len(data["results"])}</div><div class="label">Unique Signals</div></div>', unsafe_allow_html=True)
    m3.markdown(f'<div class="metric-card"><div class="value">{sum(r["count"] for r in data["results"])}</div><div class="label">Total Hits</div></div>', unsafe_allow_html=True)
    m4.markdown(f'<div class="metric-card"><div class="value risk-{data["risk_level"]}">{data["risk_level"]}</div><div class="label">Risk Level</div></div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    col_left, col_right = st.columns([2, 1])
    
    with col_right:
        st.subheader("Classification Info")
        st.markdown(f"**Detected Type:** {data['doc_type']}")
        st.markdown("**Top TF-IDF Topics:**")
        for t in data['tfidf'][:8]:
            st.markdown(f"- `{t['term']}` ({t['score']:.2f})")
            
        if data['ai']:
            st.markdown("---")
            st.subheader("🤖 AI Insights")
            st.info(data['ai'].get('summary', ''))
            st.warning(f"**Primary Risk:** {data['ai'].get('primary_risk', '')}")
            st.success(f"**Action:** {data['ai'].get('recommendation', '')}")
            if data['ai'].get('sensitive_data_types'):
                st.markdown("**Data Types Found:**")
                st.write(", ".join(data['ai']['sensitive_data_types']))

    with col_left:
        st.subheader("Extracted Signals & Context")
        
        if not data['results']:
            st.success("No sensitive signals detected in this document.")
        else:
            for kw_r in data['results']:
                cat = kw_r['category']
                st.markdown(f"**{kw_r['keyword']}** <span class='badge badge-{cat}'>{cat}</span> — **{kw_r['count']}** hit(s)", unsafe_allow_html=True)
                for ctx in kw_r['contexts'][:3]:
                    disp = ctx.replace('>>>', '<span class="hit">').replace('<<<', '</span>')
                    st.markdown(f'<div class="context-snippet">{disp}</div>', unsafe_allow_html=True)
                if len(kw_r['contexts']) > 3:
                    with st.expander(f"View remaining {kw_r['count'] - 3} hits"):
                        for ctx in kw_r['contexts'][3:]:
                            disp = ctx.replace('>>>', '<span class="hit">').replace('<<<', '</span>')
                            st.markdown(f'<div class="context-snippet">{disp}</div>', unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# View Mode: Print Mode (with Word Export)
# ---------------------------------------------------------------------------
elif st.session_state.view_mode == 'print':
    fname = st.session_state.selected_file
    gname = st.session_state.selected_group
    data = st.session_state.processed_files[gname][fname]
    
    st.markdown("## Print & Export Report")
    c1, c2 = st.columns([1, 4])
    with c1:
        if st.button("← Back"):
            st.session_state.view_mode = 'list'
            st.rerun()
    with c2:
        st.markdown("**Configure Report Content:**")
        chk1, chk2, chk3, chk4, chk5 = st.columns(5)
        inc_pii = chk1.checkbox("Include PII", value=True)
        inc_ner = chk2.checkbox("Include NER", value=True)
        inc_ai = chk3.checkbox("Include AI Analysis", value=True)
        inc_tfidf = chk4.checkbox("Include TF-IDF", value=True)
        inc_ctx = chk5.checkbox("Include Context Lines", value=True)
        
    st.markdown("---")
    
    # ---------------- WORD DOCX GENERATION ----------------
    doc = docx.Document()
    doc.add_heading(f"Security Report: {fname}", 0)
    doc.add_paragraph(f"Author: {data['metadata']['author']} | Pages: {data['metadata']['page_count']}")
    doc.add_paragraph(f"Classification: {data['doc_type']}")
    doc.add_paragraph(f"Overall Risk Level: {data['risk_level']}")
    
    if inc_ai and data['ai']:
        doc.add_heading("AI Executive Summary", level=1)
        doc.add_paragraph(data['ai'].get('summary', ''))
        doc.add_paragraph(f"Risk: {data['ai'].get('primary_risk', '')}")
        doc.add_paragraph(f"Recommendation: {data['ai'].get('recommendation', '')}")
        
    if inc_tfidf and data['tfidf']:
        doc.add_heading("Top TF-IDF Topic Signals", level=1)
        doc.add_paragraph(", ".join([f"{t['term']} ({t['score']:.2f})" for t in data['tfidf'][:10]]))
        
    doc.add_heading("Keyword Analysis", level=1)
    for kw_r in data['results']:
        is_pii = kw_r.get('is_pii', False)
        is_ner = kw_r.get('is_ner', False)
        
        if is_pii and not inc_pii: continue
        if is_ner and not inc_ner: continue
        
        p = doc.add_paragraph()
        p.add_run(f"{kw_r['keyword']} ").bold = True
        p.add_run(f"({kw_r['category'].upper()}): {kw_r['count']} occurrence(s)")
        
        if inc_ctx:
            for ctx in kw_r['contexts']:
                clean_ctx = ctx.replace('>>>', '').replace('<<<', '')
                doc.add_paragraph(f"  > ... {clean_ctx} ...", style='Quote')

    bio = io.BytesIO()
    doc.save(bio)
    
    st.download_button(
        label="📄 Download as Word Document (.docx)",
        data=bio.getvalue(),
        file_name=f"SecurityReport_{fname}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )
    
    # ---------------- VISUAL PREVIEW ----------------
    st.markdown(f"### Report Preview")
    st.markdown(f"**Classification:** {data['doc_type']} | **Risk Level:** {data['risk_level']}")
    
    if inc_ai and data['ai']:
        st.markdown("#### AI Executive Summary")
        st.info(data['ai'].get('summary', ''))
    
    if inc_tfidf and data['tfidf']:
        st.markdown("#### Top TF-IDF Topic Signals")
        st.markdown(", ".join([f"`{t['term']}`" for t in data['tfidf'][:10]]))
        
    st.markdown("#### Keyword Analysis")
    for kw_r in data['results']:
        is_pii = kw_r.get('is_pii', False)
        is_ner = kw_r.get('is_ner', False)
        if is_pii and not inc_pii: continue
        if is_ner and not inc_ner: continue
        
        st.markdown(f"- **{kw_r['keyword']}** ({kw_r['category'].upper()}): {kw_r['count']} occurrence(s)")
        if inc_ctx:
            for ctx in kw_r['contexts']:
                clean_ctx = ctx.replace('>>>', '**').replace('<<<', '**')
                st.markdown(f"  > *... {clean_ctx} ...*")
